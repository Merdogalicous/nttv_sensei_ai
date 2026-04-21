from __future__ import annotations

import importlib.util
import logging
import os
from abc import ABC, abstractmethod
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional


LOGGER = logging.getLogger(__name__)

TEXT_FILE_EXTENSIONS = {".txt", ".md"}
SUPPORTED_INGEST_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentParseError(RuntimeError):
    """Raised when a file cannot be parsed into the normalized document format."""


class DocumentParseSkipped(RuntimeError):
    """Raised when a file is intentionally skipped by configuration."""


@dataclass(frozen=True)
class ParserSettings:
    enable_pdf_ingest: bool
    pdf_parser: str
    pdf_parse_max_pages: Optional[int]
    pdf_fail_open: bool
    docling_available: bool

    @classmethod
    def from_env(cls) -> "ParserSettings":
        docling_is_available = is_docling_available()
        pdf_parser = (os.getenv("PDF_PARSER") or "docling").strip().lower() or "docling"
        enable_pdf_ingest = _resolve_pdf_ingest_enabled(pdf_parser, docling_is_available)

        raw_max_pages = (os.getenv("PDF_PARSE_MAX_PAGES") or "").strip()
        pdf_parse_max_pages = int(raw_max_pages) if raw_max_pages else None
        if pdf_parse_max_pages is not None and pdf_parse_max_pages <= 0:
            raise ValueError("PDF_PARSE_MAX_PAGES must be a positive integer when set.")

        return cls(
            enable_pdf_ingest=enable_pdf_ingest,
            pdf_parser=pdf_parser,
            pdf_parse_max_pages=pdf_parse_max_pages,
            pdf_fail_open=_parse_bool_env("PDF_FAIL_OPEN", True),
            docling_available=docling_is_available,
        )


@dataclass
class StructuredDocumentElement:
    source_file: str
    file_type: str
    parser: str
    page_start: Optional[int]
    page_end: Optional[int]
    heading_path: list[str]
    element_type: str
    text: str
    raw_metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class StructuredDocument:
    parser_name: str
    source_path: str
    file_type: str
    elements: list[StructuredDocumentElement]
    raw_metadata: dict[str, Any] = field(default_factory=dict)


class BaseDocumentParser(ABC):
    parser_name = "base"
    supported_extensions: tuple[str, ...] = ()

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.supported_extensions

    @abstractmethod
    def parse_file(self, path: Path) -> StructuredDocument:
        raise NotImplementedError


class PlainTextDocumentParser(BaseDocumentParser):
    parser_name = "plain_text"
    supported_extensions = (".txt", ".md")

    def parse_file(self, path: Path) -> StructuredDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        file_type = path.suffix.lower().lstrip(".")
        return StructuredDocument(
            parser_name=self.parser_name,
            source_path=str(path),
            file_type=file_type,
            elements=[
                StructuredDocumentElement(
                    source_file=str(path),
                    file_type=file_type,
                    parser=self.parser_name,
                    page_start=None,
                    page_end=None,
                    heading_path=[],
                    element_type="document",
                    text=text,
                    raw_metadata={"character_count": len(text)},
                )
            ],
            raw_metadata={"character_count": len(text)},
        )


class DocxDocumentParser(BaseDocumentParser):
    parser_name = "python_docx"
    supported_extensions = (".docx",)

    def parse_file(self, path: Path) -> StructuredDocument:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise DocumentParseError(
                "python-docx is not installed; cannot parse DOCX files."
            ) from exc

        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if (p.text or "").strip()]
        text = "\n".join(paragraphs)
        return StructuredDocument(
            parser_name=self.parser_name,
            source_path=str(path),
            file_type="docx",
            elements=[
                StructuredDocumentElement(
                    source_file=str(path),
                    file_type="docx",
                    parser=self.parser_name,
                    page_start=None,
                    page_end=None,
                    heading_path=[],
                    element_type="document",
                    text=text,
                    raw_metadata={"paragraph_count": len(paragraphs)},
                )
            ],
            raw_metadata={"paragraph_count": len(paragraphs)},
        )


class PyPdfDocumentParser(BaseDocumentParser):
    parser_name = "pypdf"
    supported_extensions = (".pdf",)

    def __init__(self, settings: ParserSettings, parser_name: Optional[str] = None):
        self.settings = settings
        if parser_name:
            self.parser_name = parser_name

    def parse_file(self, path: Path) -> StructuredDocument:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:
            raise DocumentParseError("pypdf is not installed; cannot parse PDF files.") from exc

        reader = PdfReader(str(path))
        page_limit = self.settings.pdf_parse_max_pages or len(reader.pages)
        elements: list[StructuredDocumentElement] = []

        for page_number, page in enumerate(reader.pages[:page_limit], start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue

            elements.append(
                StructuredDocumentElement(
                    source_file=str(path),
                    file_type="pdf",
                    parser=self.parser_name,
                    page_start=page_number,
                    page_end=page_number,
                    heading_path=[],
                    element_type="page",
                    text=text,
                    raw_metadata={"page_number": page_number},
                )
            )

        if not elements:
            raise DocumentParseError(f"No text could be extracted from PDF: {path}")

        return StructuredDocument(
            parser_name=self.parser_name,
            source_path=str(path),
            file_type="pdf",
            elements=elements,
            raw_metadata={
                "page_count": min(len(reader.pages), page_limit),
                "page_limited": self.settings.pdf_parse_max_pages is not None,
            },
        )


class DoclingPdfDocumentParser(BaseDocumentParser):
    parser_name = "docling"
    supported_extensions = (".pdf",)

    def __init__(self, settings: ParserSettings):
        self.settings = settings

    def parse_file(self, path: Path) -> StructuredDocument:
        if not self.settings.docling_available:
            raise DocumentParseError("Docling is not installed.")

        converter = _build_docling_converter()
        convert_kwargs: dict[str, Any] = {}
        if self.settings.pdf_parse_max_pages is not None:
            convert_kwargs["max_num_pages"] = self.settings.pdf_parse_max_pages

        result = converter.convert(str(path), **convert_kwargs)
        document = getattr(result, "document", None)
        if document is None:
            raise DocumentParseError(f"Docling returned no document for {path}")

        elements = _normalize_docling_document(document, source_file=str(path))
        if not elements:
            export_text = _safe_docling_export_text(document)
            if export_text:
                elements = [
                    StructuredDocumentElement(
                        source_file=str(path),
                        file_type="pdf",
                        parser=self.parser_name,
                        page_start=1 if _document_num_pages(document) else None,
                        page_end=min(
                            _document_num_pages(document) or 0,
                            self.settings.pdf_parse_max_pages or _document_num_pages(document) or 0,
                        )
                        or None,
                        heading_path=[],
                        element_type="document",
                        text=export_text,
                        raw_metadata={"normalization_mode": "export_to_text"},
                    )
                ]

        if not elements:
            raise DocumentParseError(f"Docling produced no text-bearing elements for {path}")

        return StructuredDocument(
            parser_name=self.parser_name,
            source_path=str(path),
            file_type="pdf",
            elements=elements,
            raw_metadata={
                "page_count": _document_num_pages(document),
                "conversion_status": _coerce_status(getattr(result, "status", None)),
                "used_max_num_pages": self.settings.pdf_parse_max_pages,
            },
        )


def is_docling_available() -> bool:
    return importlib.util.find_spec("docling") is not None


def parse_file(path: str | Path, settings: Optional[ParserSettings] = None) -> StructuredDocument:
    parsed_path = Path(path)
    settings = settings or ParserSettings.from_env()
    suffix = parsed_path.suffix.lower()

    if suffix in TEXT_FILE_EXTENSIONS:
        return PlainTextDocumentParser().parse_file(parsed_path)

    if suffix == ".docx":
        return DocxDocumentParser().parse_file(parsed_path)

    if suffix != ".pdf":
        raise DocumentParseError(f"Unsupported file type: {parsed_path.suffix.lower()}")

    if not settings.enable_pdf_ingest:
        raise DocumentParseSkipped(
            f"PDF ingest is disabled for {parsed_path.name}. "
            f"Set ENABLE_PDF_INGEST=true to parse PDFs."
        )

    if settings.pdf_parser == "pypdf":
        return PyPdfDocumentParser(settings=settings).parse_file(parsed_path)

    if settings.pdf_parser != "docling":
        raise DocumentParseError(
            f"Unsupported PDF_PARSER={settings.pdf_parser!r}. "
            "Supported values: docling, pypdf."
        )

    try:
        return DoclingPdfDocumentParser(settings=settings).parse_file(parsed_path)
    except DocumentParseError as exc:
        if not settings.pdf_fail_open:
            raise

        LOGGER.warning(
            "Docling PDF parsing failed for %s; falling back to pypdf. Reason: %s",
            parsed_path,
            exc,
        )
        return PyPdfDocumentParser(settings=settings, parser_name="pypdf-fallback").parse_file(parsed_path)
    except Exception as exc:
        if not settings.pdf_fail_open:
            raise DocumentParseError(f"Docling PDF parsing failed for {parsed_path}: {exc}") from exc

        LOGGER.warning(
            "Docling PDF parsing raised an unexpected error for %s; falling back to pypdf. Reason: %s",
            parsed_path,
            exc,
        )
        return PyPdfDocumentParser(settings=settings, parser_name="pypdf-fallback").parse_file(parsed_path)


def _build_docling_converter() -> Any:
    try:
        from docling.document_converter import DocumentConverter  # type: ignore
    except ImportError as exc:
        raise DocumentParseError("Docling is not installed.") from exc
    return DocumentConverter()


def _normalize_docling_document(document: Any, source_file: str) -> list[StructuredDocumentElement]:
    iterate_items = getattr(document, "iterate_items", None)
    if not callable(iterate_items):
        return []

    elements: list[StructuredDocumentElement] = []
    heading_stack: list[tuple[int, str]] = []

    for item, level in iterate_items():
        item_level = int(level or 0)
        while heading_stack and heading_stack[-1][0] >= item_level:
            heading_stack.pop()

        label = _docling_item_label(item)
        element_type = _docling_element_type(item, label)
        text = _docling_item_text(item, document).strip()
        page_start, page_end = _docling_item_pages(item)

        if not text:
            continue

        current_heading_path = [heading for _, heading in heading_stack]

        if _is_heading_label(label, element_type):
            current_heading_path = current_heading_path + [text]

        elements.append(
            StructuredDocumentElement(
                source_file=source_file,
                file_type="pdf",
                parser="docling",
                page_start=page_start,
                page_end=page_end,
                heading_path=current_heading_path,
                element_type="heading" if _is_heading_label(label, element_type) else element_type,
                text=text,
                raw_metadata={
                    "docling_label": label,
                    "level": item_level,
                    "self_ref": getattr(item, "self_ref", None),
                    "pages": _collect_docling_pages(item),
                },
            )
        )

        if _is_heading_label(label, element_type):
            heading_stack.append((item_level, text))

    return elements


def _docling_item_text(item: Any, document: Any) -> str:
    direct_text = getattr(item, "text", None)
    if isinstance(direct_text, str):
        return direct_text

    export_to_markdown = getattr(item, "export_to_markdown", None)
    if callable(export_to_markdown):
        try:
            markdown = export_to_markdown(doc=document)
        except TypeError:
            markdown = export_to_markdown()
        if isinstance(markdown, str):
            return markdown

    export_to_text = getattr(item, "export_to_text", None)
    if callable(export_to_text):
        try:
            value = export_to_text(doc=document)
        except TypeError:
            value = export_to_text()
        if isinstance(value, str):
            return value

    caption_text = getattr(item, "caption_text", None)
    if callable(caption_text):
        try:
            caption = caption_text(document)
        except TypeError:
            caption = caption_text()
        if isinstance(caption, str):
            return caption

    return ""


def _safe_docling_export_text(document: Any) -> str:
    export_to_text = getattr(document, "export_to_text", None)
    if not callable(export_to_text):
        return ""

    try:
        value = export_to_text()
    except Exception:
        return ""

    return value.strip() if isinstance(value, str) else ""


def _docling_item_label(item: Any) -> str:
    raw_label = getattr(item, "label", None)
    if raw_label is None:
        return item.__class__.__name__.lower()

    if hasattr(raw_label, "value"):
        raw_label = raw_label.value

    return str(raw_label).strip().lower()


def _docling_element_type(item: Any, label: str) -> str:
    raw_type = label or item.__class__.__name__.lower()
    normalized = raw_type.replace("docitemlabel.", "").replace(" ", "_")
    if "table" in normalized:
        return "table"
    if "picture" in normalized or "figure" in normalized or "image" in normalized:
        return "picture"
    if "list" in normalized:
        return "list_item"
    if "code" in normalized:
        return "code"
    if _is_heading_label(normalized, normalized):
        return "heading"
    return "paragraph"


def _is_heading_label(label: str, element_type: str) -> bool:
    value = f"{label} {element_type}".lower()
    heading_terms = ("heading", "section_header", "section", "title", "subtitle", "chapter")
    return any(term in value for term in heading_terms)


def _docling_item_pages(item: Any) -> tuple[Optional[int], Optional[int]]:
    pages = _collect_docling_pages(item)
    if not pages:
        return None, None
    return min(pages), max(pages)


def _collect_docling_pages(item: Any) -> list[int]:
    provenance = getattr(item, "prov", None)
    if provenance is None:
        return []

    pages: list[int] = []
    if not isinstance(provenance, Iterable):
        provenance = [provenance]

    for prov_item in provenance:
        page_no = getattr(prov_item, "page_no", None)
        if page_no is None and isinstance(prov_item, dict):
            page_no = prov_item.get("page_no")
        if isinstance(page_no, int):
            pages.append(page_no)

    return pages


def _document_num_pages(document: Any) -> Optional[int]:
    num_pages = getattr(document, "num_pages", None)
    if isinstance(num_pages, int):
        return num_pages
    if callable(num_pages):
        try:
            value = num_pages()
        except TypeError:
            value = None
        if isinstance(value, int):
            return value

    pages = getattr(document, "pages", None)
    if pages is not None:
        try:
            return len(pages)
        except TypeError:
            return None

    return None


def _coerce_status(status: Any) -> Optional[str]:
    if status is None:
        return None
    value = getattr(status, "value", None)
    if isinstance(value, str):
        return value
    name = getattr(status, "name", None)
    if isinstance(name, str):
        return name
    return str(status)


def _resolve_pdf_ingest_enabled(pdf_parser: str, docling_is_available: bool) -> bool:
    explicit = os.getenv("ENABLE_PDF_INGEST")
    if explicit is not None:
        return _parse_bool(explicit)

    if "PDF_PARSER" in os.environ:
        return True

    return pdf_parser == "docling" and docling_is_available


def _parse_bool_env(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return _parse_bool(value)


def _parse_bool(value: str) -> bool:
    normalized = value.strip().lower()
    if normalized in {"1", "true", "yes", "on"}:
        return True
    if normalized in {"0", "false", "no", "off"}:
        return False
    raise ValueError(f"Invalid boolean value: {value!r}")
